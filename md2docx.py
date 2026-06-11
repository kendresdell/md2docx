#!/usr/bin/env python3
"""
md2docx.py — Convert Markdown to a styled .docx document.

Supports:
- cover page
- footer with label + page number
- headings, paragraphs, lists, tables, blockquotes, code

All styling is controlled by a JSON style file.

Usage:
  python md2docx.py input.md output.docx
  python md2docx.py input.md output.docx --style my_style.json
"""

import argparse
import json
from pathlib import Path

import mistune
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── style constants (defaults; overwritten by build_style_constants) ──────────

FONT_NAME       = 'Poppins'
FONT_NAME_SERIF = 'Roboto Slab'

HEADING_STYLES = {
    1: {'size': 16, 'color': (0x00, 0x00, 0x00), 'bold': True,  'italic': False},
    2: {'size': 14, 'color': (0x1F, 0x38, 0x64), 'bold': True,  'italic': False},
    3: {'size': 12, 'color': (0x1F, 0x38, 0x64), 'bold': False, 'italic': True},
    4: {'size': 11, 'color': (0x40, 0x40, 0x40), 'bold': False, 'italic': True},
}
NORMAL      = {'size': 11, 'color': (0x21, 0x21, 0x21), 'bold': False, 'italic': False}
QUOTE       = {'size': 11, 'color': (0x55, 0x55, 0x55), 'bold': False, 'italic': True}
CODE_INLINE = {'size': 10, 'color': (0x33, 0x33, 0x33), 'bold': False, 'italic': False}
CODE_BLOCK  = {'size': 10, 'color': (0x33, 0x33, 0x33)}

QUOTE_BORDER_COLOR = 'A21773'
QUOTE_BG_COLOR = None

TABLE_HEADER_BG   = '2E3192'
TABLE_HEADER_TEXT = (0xFF, 0xFF, 0xFF)
TABLE_HEADER_SIZE = 10
TABLE_BODY_TEXT   = (0x43, 0x43, 0x43)
TABLE_BODY_SIZE   = 10
TABLE_BORDER_COLOR = 'D9D9D9'

DOC_LINE_SPACING = 1.15
DOC_MARGINS      = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
RENDER_THEMATIC_BREAKS = False

COVER_ENABLED = True
COVER_BG_COLOR = 'D0F5EE'
COVER_LOGO_PATH = Path('/home/kdresdell/Downloads/minipass.png')
COVER_LOGO_WIDTH_CM = 12
COVER_FONT = 'Poppins'
COVER_TITLE = "Plan d'Affaires"
COVER_TITLE_SIZE = 48
COVER_TITLE_COLOR = (0x1A, 0x1A, 0x1A)
COVER_TOP_SPACER_PT = 72

FOOTER_LABEL = "Confidentiel — Plan d’affaires minipass.me"
FOOTER_FONT = 'Poppins'
FOOTER_SIZE = 9
FOOTER_COLOR_HEX = '1A1A1A'


# ── style loading ─────────────────────────────────────────────────────────────

def load_style(path: str) -> dict:
    """Load and return a JSON style guide file."""
    with open(path, encoding='utf-8') as f:
        return json.load(f)


def hex_to_rgb(h: str) -> tuple:
    """Convert '#RRGGBB' (or 'RRGGBB') hex string to (R, G, B) integer tuple."""
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def build_style_constants(cfg: dict):
    """Populate module-level style constants from a loaded JSON style dict."""
    global FONT_NAME, FONT_NAME_SERIF, HEADING_STYLES, NORMAL, QUOTE, CODE_INLINE, CODE_BLOCK
    global QUOTE_BORDER_COLOR, QUOTE_BG_COLOR
    global TABLE_HEADER_BG, TABLE_HEADER_TEXT, TABLE_HEADER_SIZE
    global TABLE_BODY_TEXT, TABLE_BODY_SIZE, TABLE_BORDER_COLOR
    global DOC_LINE_SPACING, DOC_MARGINS, RENDER_THEMATIC_BREAKS
    global COVER_ENABLED, COVER_BG_COLOR, COVER_LOGO_PATH, COVER_LOGO_WIDTH_CM
    global COVER_FONT, COVER_TITLE, COVER_TITLE_SIZE, COVER_TITLE_COLOR, COVER_TOP_SPACER_PT
    global FOOTER_LABEL, FOOTER_FONT, FOOTER_SIZE, FOOTER_COLOR_HEX

    FONT_NAME       = cfg['fonts']['body']
    FONT_NAME_SERIF = cfg['fonts'].get('h1_override') or FONT_NAME

    heading_styles = {}
    for level, key in [(1, 'h1'), (2, 'h2'), (3, 'h3'), (4, 'h4')]:
        h = cfg['headings'][key]
        entry = {
            'size':   h['size'],
            'color':  hex_to_rgb(h['color']),
            'bold':   h['bold'],
            'italic': h['italic'],
        }
        if level == 1 and FONT_NAME_SERIF != FONT_NAME:
            entry['font'] = FONT_NAME_SERIF
        heading_styles[level] = entry
    HEADING_STYLES = heading_styles

    NORMAL = {
        'size':   cfg['body']['size'],
        'color':  hex_to_rgb(cfg['body']['color']),
        'bold':   False,
        'italic': False,
    }

    bq = cfg.get('blockquote', {})
    QUOTE = {
        'size':   bq.get('size', 11),
        'color':  hex_to_rgb(bq.get('color', '#555555')),
        'bold':   False,
        'italic': True,
    }
    QUOTE_BORDER_COLOR = bq.get('border_color', '#A21773').lstrip('#')
    quote_bg = bq.get('bg_color')
    QUOTE_BG_COLOR = quote_bg.lstrip('#') if quote_bg else None

    ci = cfg.get('code_inline', {})
    CODE_INLINE = {
        'size':   ci.get('size', 10),
        'color':  hex_to_rgb(ci.get('color', '#333333')),
        'bold':   False,
        'italic': False,
    }

    cb = cfg.get('code_block', {})
    CODE_BLOCK = {
        'size':  cb.get('size', 10),
        'color': hex_to_rgb(cb.get('color', '#333333')),
    }

    tbl = cfg.get('table', {})
    TABLE_HEADER_BG   = tbl.get('header_bg',   '#2E3192').lstrip('#')
    TABLE_HEADER_TEXT = hex_to_rgb(tbl.get('header_text', '#FFFFFF'))
    TABLE_HEADER_SIZE = tbl.get('header_size', 10)
    TABLE_BODY_TEXT   = hex_to_rgb(tbl.get('body_text',   '#434343'))
    TABLE_BODY_SIZE   = tbl.get('body_size',   10)
    TABLE_BORDER_COLOR = tbl.get('border_color', '#D9D9D9').lstrip('#')

    cov = cfg.get('cover', {})
    COVER_ENABLED       = cov.get('enabled',       False)
    COVER_BG_COLOR      = cov.get('bg_color',      '#FFFFFF').lstrip('#')
    COVER_LOGO_PATH     = Path(cov.get('logo_path', ''))
    COVER_LOGO_WIDTH_CM = cov.get('logo_width_cm', 10)
    COVER_FONT          = cfg['fonts'].get('cover_title', FONT_NAME)
    COVER_TITLE         = cov.get('title',         '')
    COVER_TITLE_SIZE    = cov.get('title_size',    36)
    COVER_TITLE_COLOR   = hex_to_rgb(cov.get('title_color', '#000000'))
    COVER_TOP_SPACER_PT = cov.get('top_spacer_pt', 72)

    ftr = cfg.get('footer', {})
    FOOTER_LABEL     = ftr.get('label', '')
    FOOTER_FONT      = cfg['fonts'].get('footer', FONT_NAME)
    FOOTER_SIZE      = ftr.get('size',  9)
    FOOTER_COLOR_HEX = ftr.get('color', '#000000').lstrip('#')

    doc_cfg = cfg.get('document', {})
    DOC_LINE_SPACING = doc_cfg.get('line_spacing', 1.15)
    RENDER_THEMATIC_BREAKS = doc_cfg.get('render_thematic_breaks', False)
    margins = doc_cfg.get('margins', {})
    DOC_MARGINS = {
        'top':    margins.get('top_cm',    2.5),
        'bottom': margins.get('bottom_cm', 2.5),
        'left':   margins.get('left_cm',   2.5),
        'right':  margins.get('right_cm',  2.5),
    }


# ── helpers ──────────────────────────────────────────────────────────────────

def _apply_run(run, style, bold=None, italic=None):
    """Apply font settings to a docx Run."""
    run.font.name      = style.get('font', FONT_NAME)
    run.font.size      = Pt(style['size'])
    run.font.color.rgb = RGBColor(*style['color'])
    run.font.bold      = bold   if bold   is not None else style['bold']
    run.font.italic    = italic if italic is not None else style['italic']


def _set_cell_bg(cell, hex_color):
    """Set background colour of a table cell via raw XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _add_para_border(para, side, color_hex, sz=6):
    """Add a single-line border on one side of a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    edge = OxmlElement(f'w:{side}')
    edge.set(qn('w:val'),   'single')
    edge.set(qn('w:sz'),    str(sz))
    edge.set(qn('w:space'), '4')
    edge.set(qn('w:color'), color_hex)
    pBdr.append(edge)
    pPr.append(pBdr)


def _set_para_bg(para, hex_color):
    """Set paragraph background colour via raw XML."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def _set_table_borders(table, hex_color):
    """Set all table borders to a specific color."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)

    for edge_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge = borders.find(qn(f'w:{edge_name}'))
        if edge is None:
            edge = OxmlElement(f'w:{edge_name}')
            borders.append(edge)
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), hex_color)


def _add_linebreak(para):
    """Insert a soft return (<w:br/>) inside a paragraph."""
    run = para.add_run()
    br  = OxmlElement('w:br')
    run._r.append(br)


def _add_page_numbers(section):
    """Footer: left label, right page number."""
    footer = section.footer
    ftr = footer._element

    for child in list(ftr):
        ftr.remove(child)

    sz_hhp = FOOTER_SIZE * 2
    rpr = (
        '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rFonts w:ascii="{FOOTER_FONT}" w:hAnsi="{FOOTER_FONT}"/>'
        f'<w:sz w:val="{sz_hhp}"/>'
        f'<w:color w:val="{FOOTER_COLOR_HEX}"/>'
        '</w:rPr>'
    )

    xml = (
        '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:tblPr>'
        '<w:tblW w:w="5000" w:type="pct"/>'
        '<w:tblBorders>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
        '</w:tblPr>'
        '<w:tr>'
        '<w:tc>'
        '<w:tcPr><w:tcW w:w="4250" w:type="pct"/></w:tcPr>'
        '<w:p><w:pPr><w:jc w:val="left"/></w:pPr>'
        f'<w:r>{rpr}<w:t xml:space="preserve">{FOOTER_LABEL}</w:t></w:r>'
        '</w:p></w:tc>'
        '<w:tc>'
        '<w:tcPr><w:tcW w:w="750" w:type="pct"/></w:tcPr>'
        '<w:p><w:pPr><w:jc w:val="right"/></w:pPr>'
        f'<w:r>{rpr}<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r>{rpr}<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
        f'<w:r>{rpr}<w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p></w:tc>'
        '</w:tr></w:tbl>'
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )

    for elem in etree.fromstring(f'<root>{xml}</root>'):
        ftr.append(elem)


def _add_cover_background(doc, para):
    """Insert a full-page solid rectangle behind cover content."""
    sec = doc.sections[0]
    cx = int(sec.page_width) if sec.page_width else 7560000
    cy = int(sec.page_height) if sec.page_height else 10692000

    xml = (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<w:drawing>'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="1" '
        'behindDoc="1" locked="1" layoutInCell="1" allowOverlap="0">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="1001" name="CoverBg"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        '<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="{COVER_BG_COLOR}"/></a:solidFill>'
        '<a:ln><a:noFill/></a:ln>'
        '</wps:spPr><wps:bodyPr/></wps:wsp></a:graphicData></a:graphic>'
        '</wp:anchor></w:drawing></w:r>'
    )
    para._p.append(etree.fromstring(xml))


def render_cover_page(doc):
    """Insert styled cover page (logo + title)."""
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(COVER_TOP_SPACER_PT)
    sp.paragraph_format.space_after = Pt(0)
    _add_cover_background(doc, sp)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(48)
    if COVER_LOGO_PATH.exists():
        run = logo_para.add_run()
        run.add_picture(str(COVER_LOGO_PATH), width=Cm(COVER_LOGO_WIDTH_CM))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(COVER_TITLE)
    run.font.name = COVER_FONT
    run.font.size = Pt(COVER_TITLE_SIZE)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*COVER_TITLE_COLOR)

    brk = doc.add_paragraph()
    run = brk.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ── inline renderer ──────────────────────────────────────────────────────────

def render_inline(para, children, style, bold=False, italic=False):
    """Recursively walk mistune inline AST nodes and add formatted runs."""
    for token in (children or []):
        t = token.get('type', '')

        if t == 'text':
            raw = token.get('raw', '')
            if raw:
                run = para.add_run(raw)
                _apply_run(run, style,
                           bold=(style['bold'] or bold),
                           italic=(style['italic'] or italic))

        elif t == 'softline':
            run = para.add_run(' ')
            _apply_run(run, style,
                       bold=(style['bold'] or bold),
                       italic=(style['italic'] or italic))

        elif t == 'linebreak':
            _add_linebreak(para)

        elif t == 'strong':
            render_inline(para, token.get('children', []), style,
                          bold=True, italic=italic)

        elif t == 'emphasis':
            render_inline(para, token.get('children', []), style,
                          bold=bold, italic=True)

        elif t == 'link':
            # Render link text (no click-through needed)
            render_inline(para, token.get('children', []), style,
                          bold=bold, italic=italic)

        elif t == 'image':
            # Skip image references silently
            pass

        elif t == 'codespan':
            run = para.add_run(token.get('raw', ''))
            run.font.name      = 'Courier New'
            run.font.size      = Pt(CODE_INLINE['size'])
            run.font.color.rgb = RGBColor(*NORMAL['color'])
            run.font.bold      = False
            run.font.italic    = False

        elif t in ('raw_html', 'html_inline'):
            pass  # skip

        elif t == 'escape':
            raw = token.get('raw', '')
            if raw:
                run = para.add_run(raw)
                _apply_run(run, style,
                           bold=(style['bold'] or bold),
                           italic=(style['italic'] or italic))

        else:
            # Generic fallback: raw text or recurse into children
            raw = token.get('raw', '')
            if raw:
                run = para.add_run(raw)
                _apply_run(run, style,
                           bold=(style['bold'] or bold),
                           italic=(style['italic'] or italic))
            elif token.get('children'):
                render_inline(para, token['children'], style,
                              bold=bold, italic=italic)


# ── block renderers ──────────────────────────────────────────────────────────

def render_heading(doc, level, children):
    style = HEADING_STYLES.get(level, NORMAL)

    # Skip empty headings (image-only or whitespace-only)
    text_content = ''.join(
        c.get('raw', '') for c in (children or [])
        if c.get('type') in ('text', 'escape')
    ).strip()
    image_only = all(c.get('type') == 'image' for c in (children or [])) if children else False
    if not children or (not text_content and not image_only and
                        not any(c.get('children') for c in (children or []))):
        return  # truly empty heading — skip

    before = {1: 18, 2: 14, 3: 10, 4: 8}.get(level, 8)
    para = doc.add_paragraph()
    # Apply the built-in heading style so Google Docs / Word navigation
    # recognise the paragraph as a heading (structure), then our run-level
    # formatting (color, font, size) overrides the visual appearance.
    try:
        para.style = doc.styles[f'Heading {level}']
    except KeyError:
        pass
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(4)
    render_inline(para, children, style)


def render_paragraph(doc, children):
    # Skip pure-image paragraphs (e.g. ![][image1])
    if children and all(c.get('type') == 'image' for c in children):
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(6)
    render_inline(para, children, NORMAL)


def render_blockquote(doc, children, extra_space_after_table=False):
    """Blockquote → indented italic paragraph with left border."""
    space_after = 11 if extra_space_after_table else 3
    for child in (children or []):
        ct = child.get('type', '')
        if ct in ('paragraph', 'block_text'):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Inches(0.35)
            para.paragraph_format.right_indent = Inches(0.2)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(space_after)
            if QUOTE_BG_COLOR:
                _set_para_bg(para, QUOTE_BG_COLOR)
            _add_para_border(para, 'left', QUOTE_BORDER_COLOR, sz=12)
            render_inline(para, child.get('children', []), QUOTE)
        elif ct == 'list':
            render_list(doc, child)
        else:
            # Recurse (nested blockquotes, etc.)
            render_block(doc, child)


def render_list(doc, token):
    ordered = token.get('attrs', {}).get('ordered', False)
    depth   = token.get('attrs', {}).get('depth', 0)
    for item in (token.get('children') or []):
        if item.get('type') == 'list_item':
            render_list_item(doc, item, ordered, depth)


def render_list_item(doc, item_token, ordered, depth):
    """Render one list item, handling nested lists."""
    children = item_token.get('children') or []

    # Separate inline content (from block_text / paragraph) vs. nested lists
    inline_children = None
    nested_lists = []

    for child in children:
        ct = child.get('type', '')
        if ct == 'list':
            nested_lists.append(child)
        elif ct in ('paragraph', 'block_text'):
            if inline_children is None:
                inline_children = child.get('children', [])
        elif ct == 'blank_line':
            pass
        else:
            # Directly inline tokens (shouldn't happen, but be safe)
            if inline_children is None:
                inline_children = []
            inline_children.append(child)

    if inline_children is None:
        inline_children = []

    para = doc.add_paragraph()
    try:
        para.style = doc.styles['List Number' if ordered else 'List Bullet']
    except KeyError:
        pass

    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    if depth > 0:
        para.paragraph_format.left_indent = Inches(0.25 + 0.25 * depth)

    render_inline(para, inline_children, NORMAL)

    # Ensure runs use the body font (list style might override font)
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(NORMAL['size'])

    # Recurse into nested lists
    for nested in nested_lists:
        nested_ordered = nested.get('attrs', {}).get('ordered', False)
        for nested_item in (nested.get('children') or []):
            if nested_item.get('type') == 'list_item':
                render_list_item(doc, nested_item, nested_ordered, depth + 1)


def render_table(doc, token):
    """Build a DOCX table from a mistune table token."""
    head_rows = []
    body_rows = []

    for section in (token.get('children') or []):
        st = section.get('type', '')
        if st == 'table_head':
            # mistune 3.x: table_head contains table_cell elements directly (no table_row wrapper)
            cells = [c for c in (section.get('children') or []) if c.get('type') == 'table_cell']
            if cells:
                head_rows.append(cells)
        elif st == 'table_body':
            for row in (section.get('children') or []):
                if row.get('type') == 'table_row':
                    body_rows.append(row.get('children') or [])

    all_rows = head_rows + body_rows
    if not all_rows:
        return

    num_cols = max(len(r) for r in all_rows)
    if num_cols == 0:
        return

    tbl = doc.add_table(rows=len(all_rows), cols=num_cols)
    tbl.style = 'Table Grid'
    _set_table_borders(tbl, TABLE_BORDER_COLOR)

    for r_idx, row_cells in enumerate(all_rows):
        is_header = r_idx < len(head_rows)
        for c_idx, cell_token in enumerate(row_cells):
            if c_idx >= num_cols:
                break

            cell = tbl.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.clear()

            align = (cell_token.get('attrs') or {}).get('align')
            if align == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            cell_children = cell_token.get('children') or []

            if is_header:
                _set_cell_bg(cell, TABLE_HEADER_BG)
                h_style = {'size': TABLE_HEADER_SIZE, 'color': TABLE_HEADER_TEXT,
                           'bold': True, 'italic': False}
                render_inline(para, cell_children, h_style)
            else:
                t_style = {'size': TABLE_BODY_SIZE, 'color': TABLE_BODY_TEXT,
                           'bold': False, 'italic': False}
                render_inline(para, cell_children, t_style)

    # Small gap after table
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(6)


def render_hr(doc):
    """Horizontal rule → thin bottom-bordered paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    _add_para_border(para, 'bottom', 'C0C0C0', sz=4)


def render_code_block(doc, token):
    """Code block — rendered with body font + indent.
    Many 'code blocks' in practice are indented flow diagrams, not source code."""
    raw = token.get('raw', '').rstrip()
    if not raw:
        return
    for line in raw.split('\n'):
        para = doc.add_paragraph()
        run  = para.add_run(line)
        run.font.name      = FONT_NAME
        run.font.size      = Pt(CODE_BLOCK['size'])
        run.font.color.rgb = RGBColor(*CODE_BLOCK['color'])
        para.paragraph_format.left_indent  = Inches(0.4)
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(1)


# ── top-level dispatcher ─────────────────────────────────────────────────────

def render_block(doc, token, next_type=None):
    t        = token.get('type', '')
    children = token.get('children') or []
    attrs    = token.get('attrs')   or {}

    if t == 'heading':
        render_heading(doc, attrs.get('level', 1), children)

    elif t == 'paragraph':
        render_paragraph(doc, children)

    elif t == 'block_quote':
        render_blockquote(doc, children, extra_space_after_table=(next_type == 'table'))

    elif t == 'list':
        render_list(doc, token)

    elif t == 'table':
        render_table(doc, token)

    elif t == 'thematic_break':
        if RENDER_THEMATIC_BREAKS:
            render_hr(doc)

    elif t == 'block_code':
        render_code_block(doc, token)

    elif t == 'blank_line':
        pass  # paragraph spacing handles whitespace

    elif t == 'block_html':
        pass  # skip raw HTML blocks

    else:
        # Unknown block type — recurse into children if any
        for child in children:
            render_block(doc, child)


def _next_nonblank_type(tokens, start_idx):
    """Return the next non-blank token type after start_idx."""
    j = start_idx + 1
    while j < len(tokens):
        t = tokens[j].get('type', '')
        if t != 'blank_line':
            return t
        j += 1
    return None


# ── document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    sec = doc.sections[0]
    sec.top_margin    = Cm(DOC_MARGINS['top'])
    sec.bottom_margin = Cm(DOC_MARGINS['bottom'])
    sec.left_margin   = Cm(DOC_MARGINS['left'])
    sec.right_margin  = Cm(DOC_MARGINS['right'])

    normal_style = doc.styles['Normal']
    normal_style.font.name      = FONT_NAME
    normal_style.font.size      = Pt(NORMAL['size'])
    normal_style.font.color.rgb = RGBColor(*NORMAL['color'])
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal_style.paragraph_format.line_spacing      = DOC_LINE_SPACING

    sect_pr = sec._sectPr
    title_pg = OxmlElement('w:titlePg')
    sect_pr.append(title_pg)
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), '0')
    sect_pr.append(pg_num_type)

    _add_page_numbers(sec)

    return doc


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('input',  help='Input .md file')
    parser.add_argument('output', help='Output .docx file')
    parser.add_argument('--style', '-s',
                        default=str(Path(__file__).parent / 'style_telus.json'),
                        help='JSON style guide (default: style_telus.json)')
    args = parser.parse_args()

    cfg    = load_style(args.style)
    build_style_constants(cfg)
    tokens = mistune.create_markdown(renderer='ast', plugins=['table'])(
                 Path(args.input).read_text(encoding='utf-8'))
    doc = setup_document()
    if COVER_ENABLED:
        render_cover_page(doc)
    for idx, token in enumerate(tokens):
        next_type = _next_nonblank_type(tokens, idx)
        render_block(doc, token, next_type=next_type)
    doc.save(args.output)
    print(f'Saved: {args.output}')


if __name__ == '__main__':
    main()
